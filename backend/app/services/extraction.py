"""File type router + text extractors (spec §4.1.1–4.1.2)."""
import io
from pathlib import Path

SUPPORTED_EXTENSIONS = {".pdf", ".docx", ".txt", ".md", ".png", ".jpg", ".jpeg", ".tiff"}


class ExtractionError(Exception):
    pass


def extract_text(filename: str, data: bytes) -> tuple[str, bool]:
    """Route a file to the right extractor. Returns (text, used_ocr)."""
    ext = Path(filename).suffix.lower()
    if ext == ".pdf":
        return _extract_pdf(data)
    if ext == ".docx":
        return _extract_docx(data), False
    if ext in (".txt", ".md"):
        return data.decode("utf-8", errors="replace"), False
    if ext in (".png", ".jpg", ".jpeg", ".tiff"):
        return _ocr_image(data), True
    raise ExtractionError(f"Unsupported file type: {ext}")


def _extract_pdf(data: bytes) -> tuple[str, bool]:
    import pdfplumber

    text_parts: list[str] = []
    with pdfplumber.open(io.BytesIO(data)) as pdf:
        for page in pdf.pages:
            text_parts.append(page.extract_text() or "")
    text = "\n".join(text_parts).strip()
    if len(text) > 50:
        return text, False
    # Likely a scanned PDF — fall back to OCR per page
    return _ocr_pdf(data), True


def _ocr_pdf(data: bytes) -> str:
    import pdfplumber
    import pytesseract

    text_parts: list[str] = []
    with pdfplumber.open(io.BytesIO(data)) as pdf:
        for page in pdf.pages:
            img = page.to_image(resolution=300).original
            text_parts.append(pytesseract.image_to_string(img))
    return "\n".join(text_parts).strip()


def _extract_docx(data: bytes) -> str:
    from docx import Document

    doc = Document(io.BytesIO(data))
    parts = [p.text for p in doc.paragraphs]
    for table in doc.tables:
        for row in table.rows:
            parts.append(" | ".join(cell.text for cell in row.cells))
    return "\n".join(parts).strip()


def _ocr_image(data: bytes) -> str:
    import pytesseract
    from PIL import Image

    img = Image.open(io.BytesIO(data))
    if img.mode != "RGB":
        img = img.convert("RGB")
    return pytesseract.image_to_string(img).strip()
